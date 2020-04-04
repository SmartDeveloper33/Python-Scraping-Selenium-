import os
from PyQt5 import QtCore, QtGui, QtWidgets, uic
from PyQt5.QtWidgets import QMessageBox
import subprocess
import csv
from shutil import rmtree

from time import sleep
import pdfkit
import csv
import os
import threading
import pathlib

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import NoSuchElementException

import requests, json
from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
from docxcompose.composer import Composer
from datetime import datetime
import comtypes.client
from tendo import singleton

class Dashboard(QtWidgets.QMainWindow):  
    
	isScraping = False
	isStopped = False
	driver = None
	dirpath = 'data'
	allpdfPath = 'data/allpdfs'
	endbookno = ""
	timer = 0
	def __init__(self, parent=None):
		super(Dashboard, self).__init__(parent)
		uic.loadUi('main.ui', self)
		self.setupUi()

	def setupUi(self):  

		#init
		if not os.path.exists(self.dirpath):
			os.mkdir(self.dirpath)
		if not os.path.exists(self.allpdfPath):
			os.mkdir(self.allpdfPath)

		# self.scTime1 = self.findChild(QtWidgets.QTimeEdit, 'scTime1')
		# self.scTime2 = self.findChild(QtWidgets.QTimeEdit, 'scTime2')
		self.btnStart = self.findChild(QtWidgets.QPushButton, 'btnScraping')
		self.btnStop = self.findChild(QtWidgets.QPushButton, 'btnStop')
		self.btnExit = self.findChild(QtWidgets.QPushButton, 'btnExit')
		self.btnDelete = self.findChild(QtWidgets.QPushButton, 'btnDelete')
		self.btnAddrValidation = self.findChild(QtWidgets.QPushButton, 'btnAddrValidation')
		self.btnExport = self.findChild(QtWidgets.QPushButton, 'btnExport')
		self.btnMakeLetter = self.findChild(QtWidgets.QPushButton, 'btnMakeLetter')
		self.btnMakeEnvelop = self.findChild(QtWidgets.QPushButton, 'btnMakeEnvelop')
		self.btnMakeLetter = self.findChild(QtWidgets.QPushButton, 'btnMakeLetter')
		self.btnMakePdfs = self.findChild(QtWidgets.QPushButton, 'btnMakePdfs')
		self.btnClearAll = self.findChild(QtWidgets.QPushButton, 'btnClearAll')
		self.btnOpen = self.findChild(QtWidgets.QPushButton, 'btnOpen')
		self.tblMain = self.findChild(QtWidgets.QTableWidget, 'tblList')
		self.tblAddr = self.findChild(QtWidgets.QTableWidget, 'tblAddr')
		self.consoleInput = self.findChild(QtWidgets.QPlainTextEdit, 'edtConsole')

		self.spinTime = self.findChild(QtWidgets.QSpinBox, 'spinTime')
		self.chkAutoScraping = self.findChild(QtWidgets.QCheckBox, 'chkAutoScraping')

		self.progressBar = self.findChild(QtWidgets.QProgressBar, 'progressBar')
		self.progressBar.hide()

		self.edtBookingNumber = self.findChild(QtWidgets.QLineEdit, 'edtBookingNumber')

		#EventListener
		self.btnStart.clicked.connect(self.onBtnStart)
		self.btnStop.clicked.connect(self.onBtnStop)
		# self.btnDelete.clicked.connect(self.deleteRow)
		self.btnAddrValidation.clicked.connect(self.addressValidation)
		self.btnExport.clicked.connect(self.onExport)
		self.btnMakeLetter.clicked.connect(self.makeLetter)
		self.btnMakeEnvelop.clicked.connect(self.makeEnvelop)
		self.btnMakePdfs.clicked.connect(self.makePdfs)
		self.btnClearAll.clicked.connect(self.clearAll)
		self.btnOpen.clicked.connect(self.open)
		self.btnExit.clicked.connect(QtWidgets.qApp.quit)

		#Timer
		self.mainTimer = QtCore.QTimer()
		self.mainTimer.timeout.connect(self.handleTimer)
		self.mainTimer.start(1000*60)

		# Init QSystemTrayIcon
		self.tray_icon = QtWidgets.QSystemTrayIcon(self)
		self.tray_icon.setIcon(self.style().standardIcon(QtWidgets.QStyle.SP_ComputerIcon))
 
		show_action = QtWidgets.QAction("Show", self)
		quit_action = QtWidgets.QAction("Exit", self)
		hide_action = QtWidgets.QAction("Hide", self)
		show_action.triggered.connect(self.show)
		hide_action.triggered.connect(self.hide)
		quit_action.triggered.connect(QtWidgets.qApp.quit)
		tray_menu = QtWidgets.QMenu()
		tray_menu.addAction(show_action)
		tray_menu.addAction(hide_action)
		tray_menu.addAction(quit_action)
		self.tray_icon.setContextMenu(tray_menu)
		self.tray_icon.show()

		self.resizeTable()

		try:
			sFile = open(self.dirpath+'/update.txt', 'r')
			self.endbookno = sFile.readline()
			sFile.close()
			self.edtBookingNumber.setText(self.endbookno)
		except IOError:
		    print("Can't access update.txt")
  
		QtCore.QMetaObject.connectSlotsByName(self)

	def resizeTable(self):
		header = self.tblMain.horizontalHeader()       
		header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(2, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(3, QtWidgets.QHeaderView.Stretch)
		header.setSectionResizeMode(4, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(5, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(6, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(7, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(8, QtWidgets.QHeaderView.Stretch)

		header = self.tblAddr.horizontalHeader()       
		header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(2, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(3, QtWidgets.QHeaderView.Stretch)
		header.setSectionResizeMode(4, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(5, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(6, QtWidgets.QHeaderView.ResizeToContents)
		header.setSectionResizeMode(7, QtWidgets.QHeaderView.ResizeToContents)

	def console(self, msg):
		self.consoleInput.appendPlainText(msg+"\t\t\t"+QtCore.QDateTime.currentDateTime().toString('yyyy-MM-dd hh:mm:ss'))
		newCursor = QtGui.QTextCursor(self.consoleInput.document())
		newCursor.movePosition(QtGui.QTextCursor.End)
		self.consoleInput.setTextCursor(newCursor)
		QtCore.QCoreApplication.processEvents()

	def open(self):
		subprocess.Popen(r'explorer "{}\data"'.format(os.getcwd()))

	def addRow(self, data):
		rowPosition = self.tblMain.rowCount()
		self.tblMain.insertRow(rowPosition)

		index = 0
		for item in data:
			self.tblMain.setItem(rowPosition , index, QtWidgets.QTableWidgetItem(item))
			index += 1

	def addRowInvalidAddr(self, data):
		rowPosition = self.tblAddr.rowCount()
		self.tblAddr.insertRow(rowPosition)

		index = 0
		for item in data:
			self.tblAddr.setItem(rowPosition , index, QtWidgets.QTableWidgetItem(item))
			index += 1

	def deleteRow(self):
		row = self.tblMain.currentRow()
		if row > -1:
			self.tblMain.removeRow(row)		

	def onExport(self):
		rowCount = self.tblMain.rowCount()
		if rowCount < 1:
			self.console("There is no data to export csv.")
			return
		colCount = self.tblMain.columnCount()
		myFile = open(self.dirpath+'/letterinfo.csv', 'w', newline='')
		writer = csv.writer(myFile)
		self.console('Outputing csv....')
		
		for i in range(0, rowCount):
			tmp = []
			for j in range(0, colCount):
				tmp.append(self.tblMain.item(i, j).text())
			writer.writerow(tmp)

		myFile.close()
		self.console('Done')

		command = str(pathlib.Path().absolute())+'/'+self.dirpath+'/letterinfo.csv'
		os.system(command)
	def makeLetter(self):
		rowCount = self.tblMain.rowCount()
		self.progressBar.setValue(0)
		self.progressBar.setMaximum(rowCount)
		self.progressBar.show()
		# merged_document = Document()
		composer = None
		for i in range(0, rowCount):
			document = Document('data/letterTemplate.docx')
			self.progressBar.setValue(i+1)
			first_name = self.tblMain.item(i, 1).text()
			booking_number = self.tblMain.item(i, 0).text()
			for paragraph in document.paragraphs:
				if '«FIRST_NAME»' in paragraph.text:
					inline = paragraph.runs
					for j in range(len(inline)):
						inline[j].text = inline[j].text.replace('«FIRST_NAME»', first_name.capitalize())
			if i < rowCount - 1: 
				document.add_page_break()
			# for element in document.element.body:
			# 	merged_document.element.body.append(element)
			if i == 0:
				composer = Composer(document)
			else :
				composer.append(document)
		# merged_document.save('data/letters.docx')
		file = "data/letters{}.docx".format(self.getCurrentDateTime())
		composer.save(file)
		self.console("Converting letters docx to pdf.")
		word = comtypes.client.CreateObject('Word.Application')
		doc = word.Documents.Open(os.getcwd()+"/"+file)
		doc.SaveAs("{}/data/letters{}.pdf".format(os.getcwd(),self.getCurrentDateTime()), FileFormat=17)
		doc.Close()
		word.Quit()
		os.remove(file)
		self.console("Done.")
		sleep(0.5)
		self.progressBar.hide()
		return
	
	def makeEnvelop(self):
		rowCount = self.tblMain.rowCount()
		self.progressBar.setValue(0)
		self.progressBar.setMaximum(rowCount)
		self.progressBar.show()
		# merged_document = Document()
		composer = None
		for i in range(0, rowCount):
			document = Document('data/letterTemplateEnvelope.docx')
			self.progressBar.setValue(i+1)
			first_name = self.tblMain.item(i, 1).text()
			last_name = self.tblMain.item(i, 2).text()
			booking_number = self.tblMain.item(i, 0).text()
			for paragraph in document.paragraphs:
				inline = paragraph.runs
				for j in range(len(inline)):
					if '«FIRST_NAME»' in inline[j].text:
						inline[j].text = first_name.capitalize()
						continue
					if '«LAST_NAME»' in inline[j].text:
						inline[j].text = last_name.capitalize()
						continue
					if '«ADDRESS»' in inline[j].text:
						inline[j].text = self.tblMain.item(i, 8).text()
						continue
					if '«ADDRESS2»' in inline[j].text:
						inline[j].text = ""
						continue
					if '«CITY»' in inline[j].text:
						inline[j].text = self.tblMain.item(i, 5).text()
					if '«STATE»' in inline[j].text:
						inline[j].text = self.tblMain.item(i, 6).text()
					if '«ZIPCODE»' in inline[j].text:
						inline[j].text = self.tblMain.item(i, 7).text()
			if i < rowCount - 1: 
				document.add_page_break()
			if i == 0:
				composer = Composer(document)
			else :
				composer.append(document)
			# for element in document.element.body:
			# 	merged_document.element.body.append(element)
		# merged_document.save('data/envelopes.docx')
		file = "data/envelopes{}.docx".format(self.getCurrentDateTime())
		composer.save(file)
		self.console("Converting envelopes docx to pdf.")
		word = comtypes.client.CreateObject('Word.Application')
		doc = word.Documents.Open(os.getcwd()+"/"+file)
		doc.SaveAs("{}/data/envelopes{}.pdf".format(os.getcwd(),self.getCurrentDateTime()), FileFormat=17)
		doc.Close()
		word.Quit()
		os.remove(file)
		self.console("Done.")
		sleep(0.5)
		self.progressBar.hide()
		return

	def makePdfs(self):
		rowCount = self.tblMain.rowCount()
		output = PdfFileWriter()
		self.progressBar.setValue(0)
		self.progressBar.setMaximum(rowCount)
		self.progressBar.show()
		tmp_file_handles = []
		for i in range(0, rowCount):
			self.progressBar.setValue(i+1)
			booking_number = self.tblMain.item(i, 0).text()
			try:
				tmp = open("{}/{}.pdf".format(self.allpdfPath,booking_number), "rb")
				tmp_file_handles.append(tmp)
				inputPdf = PdfFileReader(tmp)
				for i in range(inputPdf.numPages):
					output.addPage(inputPdf.getPage(i))
			except Exception as e:
				self.console("Add pdf error: booking number is {}".format(booking_number))
				print("Add pdf error: booking number is {}".format(booking_number))
		with open("data/print{}.pdf".format(self.getCurrentDateTime()), "wb") as outputStream:
			output.write(outputStream)
		for fHandle in tmp_file_handles:
			fHandle.close()

		self.progressBar.setValue(rowCount)
		sleep(0.5)
		self.progressBar.hide()
		return

	def clearAll(self):
		self.tblMain.setRowCount(0)
		self.tblAddr.setRowCount(0)
		rmtree(self.allpdfPath)
		self.consoleInput.clear()

		sFile = open(self.dirpath+'/update.txt', 'w+')
		sFile.write(self.edtBookingNumber.text())
		sFile.close()

		return

	def getCurrentDateTime(self):
		month = datetime.now().date().month
		day = datetime.now().date().day
		if month < 10:
			month = "0"+str(month)
		if day < 10:
			day = "0" + str(day)
		hour = datetime.now().time().hour
		if hour < 10:
			hour = "0" + str(hour)
		minute = datetime.now().time().minute
		if minute < 10:
			minute = "0" + str(minute)

		return "{}{}_{}{}".format(month,day, hour,minute)

	def validAddress(self, payload):		 
		headers = {"API-KEY": "TEST_h7uzZHUTRVpgTkKrHGNwdwl5Maq7A70zZ2Rq4Vdw5v8", "Content-Type": "application/json"}
		url = "https://api.shipengine.com/v1/addresses/validate"
		response = requests.post(url, headers=headers, data=json.JSONEncoder().encode(payload))

		return response.json()

	def addressValidation(self):
		QtWidgets.QApplication.setOverrideCursor(QtGui.QCursor(QtCore.Qt.WaitCursor))
		rowCount = self.tblMain.rowCount()
		index = 0
		payload = []
		result = []
		real_addr = []
		self.console('Starting address validation...')
		self.progressBar.show()
		self.progressBar.setMaximum(rowCount+10)
		self.progressBar.setValue(0)
		for i in range(0, rowCount):
			self.progressBar.setValue(i+1)
			tmp = {}
			tmp["address_line1"] = self.tblMain.item(i, 3).text()
			tmp["city_locality"] = self.tblMain.item(i, 5).text()
			tmp["state_province"] = self.tblMain.item(i, 6).text()
			tmp["postal_code"] = ""
			tmp["country_code"] = "US"
			payload.append(tmp)
			if index == 100:
				index = 0
				response = self.validAddress(payload)
				for addr in response:
					if addr['status'] == 'verified':
						result.append(True)
						real_addr.append(addr['matched_address']['address_line1'])
					else:
						result.append(False)
						real_addr.append('')
				payload = []
			index += 1
		self.progressBar.setValue(rowCount+5)
		response = self.validAddress(payload)
		for addr in response:
			if addr['status'] == 'verified':
				result.append(True)
				real_addr.append(addr['matched_address']['address_line1'])
			else:
				result.append(False)
				real_addr.append('')
		removeRows = []
		index = 0
		self.progressBar.setValue(rowCount+7)
		for isValid in result:
			if isValid:
				self.tblMain.setItem(index, 8, QtWidgets.QTableWidgetItem(real_addr[index]))
			else:
				removeRows.append(index)
			index += 1

		#remove rows
		removeRows.reverse()
		colCnt = self.tblMain.columnCount()
		for i in removeRows:
			rowPosition = self.tblAddr.rowCount()
			self.tblAddr.insertRow(rowPosition)
			for j in range(0, colCnt - 1):
				self.tblAddr.setItem(rowPosition, j, QtWidgets.QTableWidgetItem(self.tblMain.item(i, j)))
			self.tblMain.removeRow(i)
		self.resizeTable()
		QtWidgets.QApplication.restoreOverrideCursor()
		self.console('Address validation is ended.')
		self.progressBar.setValue(rowCount+10)
		sleep(1)
		self.progressBar.hide()
	def getAddress(self, addr):
		address = ['','']
		address[0] = addr
		address[1] = ''

		temp = addr.split('RD ')
		if len(temp) > 1:
			address[0] = temp[0] + 'RD'
			address[1] = temp[1]
			return address

		temp = addr.split('ST ')
		if len(temp) > 1:
			address[0] = temp[0] + 'ST'
			address[1] = temp[1]		
			return address

		temp = addr.split('APT ')
		if len(temp) > 1:
			address[0] = temp[0] + 'APT'
			address[1] = temp[1]		
			return address

		return address 

	def scraping(self):

		self.endbookno = self.edtBookingNumber.text()
		#self.consoleInput.clear()
		self.console("Scraping started!")
		driver = webdriver.Edge(executable_path = 'msedgedriver.exe')
		sleep(2)
		driver.get('http://www.lubbocksheriff.com/active-jail-roster')
		

		iframe = driver.find_element_by_xpath("//iframe[@class='iframe-class']")
		sleep(2)
		driver.switch_to.frame(iframe)
		#Order by
		Book_No = driver.find_element_by_xpath('.//form/table/tbody/tr/td/table/tbody/tr[4]/td/table/tbody/tr[1]/td[1]/a')
		

		sleep(1)
		Book_No.click()
		sleep(1)
		Book_No = driver.find_element_by_xpath('.//form/table/tbody/tr/td/table/tbody/tr[4]/td/table/tbody/tr[1]/td[1]/a')
		Book_No.click()
		sleep(1)

		
		firstBookNo = ""
		timeout = 5

		while(1):
			self.isScraping = True
			self.isEnd = False;
			for i in range(2, 12) :
				if self.isStopped == True:
					self.console('Scraping is stopped by user')
					self.isStopped = False
					self.isScraping = False
					self.isEnd = True
					break
				tempData = []
				if i > 2:
					iframe = driver.find_element_by_xpath("//iframe[@class='iframe-class']")
					driver.switch_to.frame(iframe)
					sleep(1)
				try:
				    ele = EC.presence_of_element_located((By.XPATH, './/form/table/tbody/tr/td/table/tbody/tr[4]/td/table/tbody/tr[{}]'.format(i)))
				    WebDriverWait(driver, timeout).until(ele)
				except TimeoutException:
				    self.console("Timed out waiting for page to load")
				    self.isScraping = False

				tr = driver.find_element_by_xpath('.//form/table/tbody/tr/td/table/tbody/tr[4]/td/table/tbody/tr[{}]'.format(i))
				dirName = tr.find_element_by_xpath('td[1]').text
				self.console('Booking number: {}'.format(dirName))
				if int(dirName) <= int(self.endbookno):
					self.isEnd = True
					break
				tempData.append(dirName)

				tr.find_element_by_xpath('td[7]//input').click()
				driver.switch_to.window(driver.window_handles[len(driver.window_handles) - 1])
				sleep(1)

				try:
				    ele = EC.presence_of_element_located((By.XPATH, './/frameset/frame[1]'))
				    WebDriverWait(driver, timeout).until(ele)
				except TimeoutException:
				    self.console("Timed out waiting for page to load")
				    self.isScraping = False
				frame1 = driver.find_element_by_xpath('.//frameset/frame[1]')
				driver.switch_to.frame(frame1)
				sleep(1.5)
				try:
				    ele = EC.presence_of_element_located((By.ID, 'Label1'))
				    WebDriverWait(driver, timeout).until(ele)
				except TimeoutException:
				    self.console("Timed out waiting for page to load")
				    self.isScraping = False

				names = driver.find_element_by_id('Label1').text.split(',')
				first_name = ""
				last_name = ""
				if len(names) > 1 :
					last_name = names[0].strip()
					first_name = names[1].strip().split(" ")[0]
				tempData.append(first_name)
				tempData.append(last_name)

				addr = driver.find_element_by_id('addr').text
				# validationAddr = self.validAdress(addr)
				
				# tmp = self.getAddress(addr)
				# addr = tmp[0]
				# addr2 = tmp[1]

				citystzip = driver.find_element_by_id('citystzip').text.split(',')
				city = ""
				state = ""
				zipcode = ""
				if len(citystzip) > 1:
					city = citystzip[0].strip()
					state = citystzip[1].strip().split(' ')[0]
					if len(citystzip[1].strip().split(' ')) > 1:
						zipcode = citystzip[1].strip().split(' ')[1]	
				tempData.append(addr)
				tempData.append('')
				tempData.append(city)
				tempData.append(state)
				tempData.append(zipcode)

				# if validationAddr == True:
				self.addRow(tempData)
				# else :
					# self.addRowInvalidAddr(tempData)

				# if validationAddr == True:
				html = driver.find_element_by_xpath('.//form').get_attribute('innerHTML')
				image = driver.find_element_by_tag_name('img').get_attribute('src')
				html = html.replace('../','https://apps.co.lubbock.tx.us/')

				if not os.path.exists(self.allpdfPath):
					os.mkdir(self.allpdfPath)

				self.console('Output pdf....')
				fileName = '{}/{}.pdf'.format(self.allpdfPath, dirName)
				pdfkit.from_string(html, fileName)
				self.console('Done')
				# copyfile(fileName, '{}/{}.pdf'.format(self.allpdfPath,dirName))

				try:
					driver.switch_to.default_content()
				except Exception as e:
					self.console(e)
					self.isScraping = False

				frame2 = driver.find_element_by_xpath('.//frameset/frame[2]')
				driver.switch_to.frame(frame2)
				sleep(2)	
				driver.find_element_by_id('close').click()
				driver.switch_to.window(driver.window_handles[0])

				if firstBookNo == "":
					firstBookNo = dirName
					# sFile = open(self.dirpath+'/update.txt', 'w+')
					# sFile.write(firstBookNo)
					self.edtBookingNumber.setText(firstBookNo)
					# sFile.close()

				if self.endbookno == "":
					self.endbookno = '{}'.format(int(firstBookNo) - 30)
					self.console("EndBookNo: {}".format(self.endbookno) )
				
				sleep(1)

			if self.isEnd:
				self.console("Scraping is ended.")
				self.isScraping = False
				break
			self.console('Moving next page')
			iframe = driver.find_element_by_xpath("//iframe[@class='iframe-class']")
			driver.switch_to.frame(iframe)
			sleep(1)
			nextPage = driver.find_element_by_xpath('.//form/table/tbody/tr/td/table/tbody/tr[4]/td/table/tbody/tr[12]/td/a[text()="Next"]'.format(i))
			nextPage.click()

		# self.resizeTable()
		sleep(2)
		driver.close()
    #slots
	def onBtnStart(self):
		if self.isScraping == True:
			self.console("In scraping..")
			return
		t = threading.Thread(target=self.scraping)
		t.start()

	def onBtnStop(self):
		if self.isScraping == True:
			self.console('Stopping...')
			self.isStopped = True
		else :
			self.console('There is no scraping...')

	def handleTimer(self):
		if self.chkAutoScraping.isChecked() == False:
			return

		if self.isScraping == True:
			return

		self.timer += 1
		
		if self.timer >= self.spinTime.value():
			self.console("Schedule Scraping...")
			t = threading.Thread(target=self.scraping)
			t.start()
			self.timer = 0

	def closeEvent(self, event):
		event.ignore()
		self.hide()
		self.tray_icon.showMessage(
			"Scraping Tool",
			"Application was minimized to Tray",
			QtWidgets.QSystemTrayIcon.Information,
			2000
		)

if __name__ == "__main__":  
	me = singleton.SingleInstance()
	import sys  
	app = QtWidgets.QApplication(sys.argv)  
	home = Dashboard()
	home.showMaximized() 
	sys.exit(app.exec())